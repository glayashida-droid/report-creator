from pathlib import Path
from typing import List, Optional, Set

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from src.generators.word_engine import (
    FONT_EAST,
    SIZE_BODY,
    SIZE_PHOTO_TITLE,
    WordGenerator,
    _WIDTHS_COVER_INFO,
    _WIDTHS_SAMPLE_LIST,
    _WIDTHS_SUMMARY,
)
from src.io.test_photos import test_dir_key as leg_test_dir_key
from src.models.project_state import (
    DataTableRef,
    ProjectState,
    TestLeg,
    TestNode,
    TestSample,
    TestResult,
)


def _east_font(run) -> Optional[str]:
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        return None
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        return None
    return rf.get(qn("w:eastAsia"))


def _grid_widths(table) -> List[int]:
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        return []
    return [int(gc.get(qn("w:w"))) for gc in grid.findall(qn("w:gridCol"))]


def _tbl_border_vals(table) -> Set[str]:
    tblPr = table._tbl.tblPr
    if tblPr is None:
        return set()
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        return set()
    return {child.get(qn("w:val")) for child in borders}


def test_word_engine_zh_template(tmp_path=None):
    state = ProjectState(
        project_id="A2260613686101",
        applicant_name="Test Client Co.",
        sample_name="Engine Control Unit",
        sample_receive_date="2026-08-01",
        test_start_date="2026-08-02",
        test_end_date="2026-08-10",
        application_fields={
            "申请单号": "A22606136861",
            "申请公司": "Test Client Co.",
            "申请公司地址": "Shanghai",
            "样品名称": "Engine Control Unit",
            "零件号": "P-001",
            "送样数量": "2",
        },
    )
    node1 = TestNode(
        test_name="盐雾试验",
        standard_id="VW 80000",
        standard_chapter="4.1",
        standard_desc="在35度环境下喷洒盐水...",
        equipment_name="盐雾试验箱",
        evaluation_req="表面无腐蚀",
        env_condition="(25±5)°C (50±25)%Rh",
        samples=[
            TestSample(sample_id="A01", result=TestResult.PASS, result_desc="无腐蚀"),
            TestSample(sample_id="A02", result=TestResult.PASS, result_desc="无腐蚀"),
        ],
    )
    leg1 = TestLeg(leg_id="L1", leg_name="Leg 1")
    leg1.nodes.append(node1)
    state.legs.append(leg1)

    template_path = Path("templates/template_zh.docx")
    if not template_path.exists():
        template_path = Path("templates/template_raw.docx")
    out_dir = Path(tmp_path) if tmp_path else Path(".scratch")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "output_test.docx"

    engine = WordGenerator(str(template_path))
    engine.generate(state, str(out_path), project_path=None, report_language="中文")

    assert out_path.exists()

    doc = Document(str(out_path))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "盐雾试验" in texts
    assert "检测环境条件" in texts
    assert "(25±5)°C (50±25)%Rh" in texts
    assert "{{试验明细}}" not in texts
    assert "{{样品信息表}}" not in texts
    assert "No." in texts
    # header report number filled
    hdr = "\n".join(c.text for t in doc.sections[0].header.tables for r in t.rows for c in r.cells)
    assert "{{报告编号}}" not in hdr
    assert "A226061368610100001C" in hdr
    assert WordGenerator.default_report_no(state, "英文") == "A226061368610100001"
    assert WordGenerator.default_report_no(state, "中英文") == "A226061368610100001E"

    # Formatting contract vs CTI golden reports
    cover = None
    sample_list = None
    for t in doc.tables:
        first = (t.rows[0].cells[0].text or "").strip()
        if first.startswith("样品名称") or first == "样品名称":
            cover = t
        if first == "序号" and len(t.columns) == 5:
            second = (t.rows[0].cells[1].text or "").strip()
            if second == "样品名称":
                sample_list = t
    assert cover is not None
    assert sample_list is not None
    assert _grid_widths(cover) == list(_WIDTHS_COVER_INFO)
    assert _tbl_border_vals(cover) == {"nil"}
    assert _grid_widths(sample_list) == list(_WIDTHS_SAMPLE_LIST)

    cover_run = cover.rows[0].cells[0].paragraphs[0].runs[0]
    assert cover_run.bold is not True
    assert _east_font(cover_run) == FONT_EAST
    assert cover_run.font.size is not None
    assert abs(cover_run.font.size.pt - SIZE_BODY) < 0.1

    title_para = next(p for p in doc.paragraphs if p.text.startswith("1. 检测项目："))
    assert all(r.bold is not True for r in title_para.runs)
    assert title_para.paragraph_format.space_before is not None
    assert title_para.paragraph_format.space_before.pt == 0
    assert abs((title_para.paragraph_format.line_spacing or 0) - (276 / 240)) < 0.01
    snap = title_para._element.find(qn("w:pPr"))
    assert snap is not None
    snap_el = snap.find(qn("w:snapToGrid"))
    assert snap_el is not None and snap_el.get(qn("w:val")) == "0"

    photo_para = next(p for p in doc.paragraphs if p.text.strip() == "检测照片")
    assert all(r.bold is not True for r in photo_para.runs)
    assert photo_para.runs
    assert abs(photo_para.runs[0].font.size.pt - SIZE_PHOTO_TITLE) < 0.1

    # Cover row height + page breaks before test item and photos
    tr = cover.rows[0]._tr
    trPr = tr.find(qn("w:trPr"))
    assert trPr is not None
    trH = trPr.find(qn("w:trHeight"))
    assert trH is not None
    assert int(trH.get(qn("w:val"))) == 200

    # Rows must not split across pages (Word: 允许跨页断行 = off)
    sample_result = None
    for t in doc.tables:
        first = (t.rows[0].cells[0].text or "").strip()
        if first == "样品编号" and len(t.columns) == 3:
            sample_result = t
            break
    assert sample_result is not None
    for tbl in (cover, sample_list, sample_result):
        for row in tbl.rows:
            trPr = row._tr.find(qn("w:trPr"))
            assert trPr is not None and trPr.find(qn("w:cantSplit")) is not None
    # Sample result header repeats on each page
    hdr_pr = sample_result.rows[0]._tr.find(qn("w:trPr"))
    assert hdr_pr is not None and hdr_pr.find(qn("w:tblHeader")) is not None
    data_pr = sample_result.rows[1]._tr.find(qn("w:trPr"))
    assert data_pr is None or data_pr.find(qn("w:tblHeader")) is None

    def _page_break_before(para) -> bool:
        # Walk previous siblings for a page-break paragraph
        el = para._element.getprevious()
        while el is not None:
            tag = el.tag.split("}")[-1]
            if tag == "p":
                for br in el.findall(".//" + qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        return True
                # stop if previous non-empty text paragraph without break
                texts = [t.text or "" for t in el.findall(".//" + qn("w:t"))]
                if "".join(texts).strip():
                    return False
            el = el.getprevious()
        return False

    assert _page_break_before(title_para)
    assert _page_break_before(photo_para)

    print(f"Generated test report to {out_path}")


def test_condition_image_width_caps_small_pngs():
    from PIL import Image
    import io

    buf = io.BytesIO()
    Image.new("RGB", (300, 150), color=(200, 200, 200)).save(buf, format="PNG")
    width = WordGenerator._condition_image_width_in(buf.getvalue())
    # 300px / 150dpi = 2.0in, not stretched to 5in
    assert abs(width - 2.0) < 0.05

    buf2 = io.BytesIO()
    Image.new("RGB", (2000, 800), color=(100, 100, 100)).save(buf2, format="PNG")
    width2 = WordGenerator._condition_image_width_in(buf2.getvalue())
    assert abs(width2 - 5.0) < 0.05


def test_condition_blocks_titled_and_indented(tmp_path):
    from docx import Document

    from src.models.project_state import TestStandard

    template = tmp_path / "t.docx"
    Document().save(template)
    node = TestNode(test_name="组合")
    node.apply_standards(
        [
            TestStandard(
                standard_id="S1",
                chapter="1",
                test_name="机械冲击",
                standard_desc="冲击条件第一段。\n冲击条件第二段。",
            ),
            TestStandard(
                standard_id="S2",
                chapter="2",
                test_name="防尘实验",
                standard_desc="防尘条件正文。",
            ),
        ]
    )
    gen = WordGenerator(str(template))
    doc = Document()
    anchor = doc.add_paragraph("ANCHOR")
    gen._insert_condition_blocks(doc, anchor, node)
    texts = [p.text for p in doc.paragraphs if p.text != "ANCHOR"]
    assert texts == [
        "机械冲击",
        "  冲击条件第一段。",
        "  冲击条件第二段。",
        "防尘实验",
        "  防尘条件正文。",
    ]


def test_prepare_embed_stream_downscales_to_330ppi():
    from PIL import Image
    import io
    import os

    from src.generators.word_engine import EMBED_PPI, PHOTO_WIDTH_IN

    # Oversized photo-like noise (> 330ppi @ 2.95in ≈ 974px)
    w, h = 3000, 2250
    im = Image.frombytes("RGB", (w, h), os.urandom(w * h * 3))
    src = io.BytesIO()
    im.save(src, format="JPEG", quality=95)
    raw = src.getvalue()
    assert len(raw) > 200_000

    out = WordGenerator._prepare_embed_stream(raw, PHOTO_WIDTH_IN)
    data = out.getvalue()
    with Image.open(io.BytesIO(data)) as im2:
        assert im2.format == "JPEG"
        target = WordGenerator._target_embed_px(PHOTO_WIDTH_IN)
        assert im2.size[0] == target
        assert abs(im2.size[0] / PHOTO_WIDTH_IN - EMBED_PPI) < 1.0
        assert abs(im2.size[1] / im2.size[0] - h / w) < 0.02
    assert len(data) < len(raw) / 2
    # q=96 textured content should not collapse under typical golden phone crops
    assert len(data) >= 80_000


def test_prepare_embed_stream_keeps_small_jpeg():
    from PIL import Image
    import io

    from src.generators.word_engine import PHOTO_WIDTH_IN

    src = io.BytesIO()
    Image.new("RGB", (800, 600), color=(10, 20, 30)).save(src, format="JPEG", quality=90)
    raw = src.getvalue()
    out = WordGenerator._prepare_embed_stream(raw, PHOTO_WIDTH_IN)
    assert out.getvalue() == raw


def test_word_engine_data_table_merges(tmp_path):
    from openpyxl import Workbook

    project = Path(tmp_path) / "proj"
    attach = project / "3.测试组" / leg_test_dir_key("L1", "冲击") / "数据表附件"
    attach.mkdir(parents=True)
    xlsx = attach / "resist.xlsx"
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "样品编号 Sample No."
    ws.merge_cells("A1:A4")
    ws["B1"] = "试验前"
    ws["E1"] = "试验后"
    ws.merge_cells("B1:D1")
    ws.merge_cells("E1:G1")
    ws["B2"] = "Before test"
    ws["E2"] = "After test"
    ws.merge_cells("B2:D2")
    ws.merge_cells("E2:G2")
    ws["B3"] = "桥路电阻"
    ws["C3"] = "短路电阻"
    ws["D3"] = "绝缘电阻"
    ws["E3"] = "桥路电阻"
    ws["F3"] = "短路电阻"
    ws["G3"] = "绝缘电阻"
    ws["B4"] = "(Ω)"
    ws["C4"] = "(Ω)"
    ws["D4"] = "(MΩ)"
    ws["E4"] = "(Ω)"
    ws["F4"] = "(Ω)"
    ws["G4"] = "(MΩ)"
    ws["A5"] = "TP-1"
    ws["B5"] = "1.92"
    wb.save(xlsx)
    wb.close()

    state = ProjectState(
        project_id="T1",
        application_fields={"申请单号": "T1", "样品名称": "S"},
    )
    node = TestNode(
        test_name="冲击",
        samples=[TestSample(sample_id="A01", result=TestResult.PASS)],
        data_tables=[
            DataTableRef(title="resist.xlsx", relative_path=str(xlsx.relative_to(project)))
        ],
    )
    leg = TestLeg(leg_id="L1", leg_name="L1")
    leg.nodes.append(node)
    state.legs.append(leg)

    template_path = Path("templates/template_zh.docx")
    if not template_path.exists():
        template_path = Path("templates/template_raw.docx")
    out_path = Path(tmp_path) / "merge.docx"
    WordGenerator(str(template_path)).generate(
        state, str(out_path), project_path=str(project), report_language="中文"
    )

    doc = Document(str(out_path))
    data_tbl = None
    for t in doc.tables:
        if (t.rows[0].cells[0].text or "").startswith("样品编号") and len(t.columns) == 7:
            data_tbl = t
            break
    assert data_tbl is not None
    widths = _grid_widths(data_tbl)
    assert len(widths) == 7
    assert len(set(widths)) > 1  # not equal-split

    # B1:D1 and E1:G1 should be merged → unique cells in row 0 fewer than 7
    seen = set()
    unique = []
    for cell in data_tbl.rows[0].cells:
        if id(cell._tc) in seen:
            continue
        seen.add(id(cell._tc))
        unique.append(cell.text)
        # No stacked empty paragraphs from merge
        assert len(cell.paragraphs) == 1, [p.text for p in cell.paragraphs]
    assert len(unique) == 3
    assert "试验前" in unique[1]
    assert "试验后" in unique[2]

    # Header band (rows 0-3) marked tblHeader for repeat-on-each-page
    for ri in range(4):
        trPr = data_tbl.rows[ri]._tr.find(qn("w:trPr"))
        assert trPr is not None and trPr.find(qn("w:tblHeader")) is not None
    # Data row is not a header
    trPr5 = data_tbl.rows[4]._tr.find(qn("w:trPr"))
    assert trPr5 is None or trPr5.find(qn("w:tblHeader")) is None

    data_label = next(p for p in doc.paragraphs if p.text.strip() == "试验数据")
    assert data_label.alignment == WD_ALIGN_PARAGRAPH.CENTER
    title_para = next(p for p in doc.paragraphs if p.text.strip() == "resist")
    assert title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_word_engine_data_table_two_row_header_by_sample_id(tmp_path):
    from openpyxl import Workbook

    project = Path(tmp_path) / "proj"
    attach = project / "3.测试组" / leg_test_dir_key("L1", "冲击") / "数据表附件"
    attach.mkdir(parents=True)
    xlsx = attach / "after.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.merge_cells("B1:D1")
    ws["B1"] = "试验后 after test"
    ws["B2"] = "桥路电阻"
    ws["C2"] = "短路电阻"
    ws["D2"] = "绝缘电阻"
    ws["A3"] = "A22607480801-A01"
    ws["B3"] = "1.2"
    wb.save(xlsx)
    wb.close()

    state = ProjectState(
        project_id="T1",
        application_fields={"申请单号": "T1", "样品名称": "S"},
    )
    node = TestNode(
        test_name="冲击",
        samples=[TestSample(sample_id="A22607480801-A01", result=TestResult.PASS)],
        data_tables=[
            DataTableRef(title="after.xlsx", relative_path=str(xlsx.relative_to(project)))
        ],
    )
    leg = TestLeg(leg_id="L1", leg_name="L1")
    leg.nodes.append(node)
    state.legs.append(leg)

    template_path = Path("templates/template_zh.docx")
    if not template_path.exists():
        template_path = Path("templates/template_raw.docx")
    out_path = Path(tmp_path) / "two_row_header.docx"
    WordGenerator(str(template_path)).generate(
        state, str(out_path), project_path=str(project), report_language="中文"
    )

    doc = Document(str(out_path))
    data_tbl = None
    for t in doc.tables:
        if (t.rows[0].cells[1].text or "").startswith("试验后") and len(t.columns) == 4:
            data_tbl = t
            break
    assert data_tbl is not None
    for ri in range(2):
        trPr = data_tbl.rows[ri]._tr.find(qn("w:trPr"))
        assert trPr is not None and trPr.find(qn("w:tblHeader")) is not None
    trPr_data = data_tbl.rows[2]._tr.find(qn("w:trPr"))
    assert trPr_data is None or trPr_data.find(qn("w:tblHeader")) is None


def test_summary_table_testing_period_column(tmp_path):
    state = ProjectState(
        project_id="A2260613686101",
        application_fields={"申请单号": "A22606136861", "样品名称": "ECU"},
    )
    node = TestNode(
        test_name="振动试验",
        start_date="2026-07-03",
        end_date="2026-07-23",
        samples=[TestSample(sample_id="A01", result=TestResult.PASS)],
    )
    leg = TestLeg(leg_id="L1", leg_name="Leg 1", nodes=[node])
    state.legs.append(leg)

    templates = {
        "中文": Path("templates/template_zh.docx"),
        "英文": Path("templates/template_en.docx"),
        "中英文": Path("templates/template_ze.docx"),
    }

    for lang, header, period in [
        ("中文", "检测时间", "2026.07.03~2026.07.23"),
        ("英文", "Testing Period", "2026.07.03~2026.07.23"),
        ("中英文", "检测时间\nTesting Period", "2026.07.03~2026.07.23"),
    ]:
        template_path = templates[lang]
        if not template_path.exists():
            template_path = Path("templates/template_raw.docx")
        out_path = Path(tmp_path) / f"summary_{lang}.docx"
        WordGenerator(str(template_path)).generate(
            state, str(out_path), report_language=lang
        )
        doc = Document(str(out_path))
        summary = None
        for t in doc.tables:
            headers = [(c.text or "").strip() for c in t.rows[0].cells]
            if len(headers) != 6:
                continue
            joined = " ".join(headers)
            if ("检测项目" in joined or "Test Item" in joined) and (
                "检测时间" in joined or "Testing Period" in joined
            ):
                summary = t
                break
        assert summary is not None, lang
        assert len(summary.columns) == 6, lang
        assert _grid_widths(summary) == list(_WIDTHS_SUMMARY)
        headers = [(c.text or "").strip() for c in summary.rows[0].cells]
        assert headers[3] == header, lang
        assert summary.rows[1].cells[3].text.strip() == period, lang


def test_fmt_period_zero_pads_month():
    assert WordGenerator._fmt_period("2026-01-05", "2026-09-07") == "2026.01.05~2026.09.07"


def test_custom_report_no_in_document(tmp_path):
    state = ProjectState(
        project_id="A2260613686101",
        applicant_name="Test Client Co.",
        sample_name="Engine Control Unit",
        sample_receive_date="2026-08-01",
        test_start_date="2026-08-02",
        test_end_date="2026-08-10",
        application_fields={"申请单号": "A22606136861", "样品名称": "Engine Control Unit"},
    )
    leg = TestLeg(leg_id="L1", leg_name="Leg 1")
    leg.nodes.append(
        TestNode(
            test_name="盐雾试验",
            samples=[TestSample(sample_id="A01", result=TestResult.PASS)],
        )
    )
    state.legs.append(leg)

    template_path = Path("templates/template_zh.docx")
    if not template_path.exists():
        template_path = Path("templates/template_raw.docx")
    out_path = Path(tmp_path) / "custom_no.docx"
    custom_no = "A226061368610100002C"
    WordGenerator(str(template_path)).generate(
        state,
        str(out_path),
        project_path=None,
        report_language="中文",
        report_no=custom_no,
    )

    doc = Document(str(out_path))
    hdr = "\n".join(
        c.text for t in doc.sections[0].header.tables for r in t.rows for c in r.cells
    )
    assert custom_no in hdr
    assert "A226061368610100001C" not in hdr


def test_report_filename_stem_and_duplicate_path(tmp_path):
    assert WordGenerator.report_filename_stem("A226061368610100001C") == "A226061368610100001C"
    assert WordGenerator.report_filename_stem("bad/name") == "bad_name"

    folder = Path(tmp_path)
    (folder / "A226061368610100001C.docx").write_bytes(b"x")
    alt = WordGenerator.next_duplicate_report_path(folder, "A226061368610100001C")
    assert alt.name == "A226061368610100001C-2.docx"

    alt.write_bytes(b"y")
    alt2 = WordGenerator.next_duplicate_report_path(folder, "A226061368610100001C")
    assert alt2.name == "A226061368610100001C-3.docx"


if __name__ == "__main__":
    test_word_engine_zh_template()
