"""Same-as-applicant report-title EN must not fall back to Chinese."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook

from application_parser.excel_parser import parse_application_sheet1
from application_parser.field_extract_applicant import resolve_same_as_applicant
from src.generators.word_engine import WordGenerator
from src.models.project_state import ProjectState


def test_resolve_same_as_applicant_en_side_no_cn_fallback():
    cn = "奥托立夫（上海）汽车安全系统研发有限公司"
    en = "Autoliv (Shanghai) Automotive Safety System R&D Co., Ltd."
    assert (
        resolve_same_as_applicant(
            "同申请公司 Same as Applicant",
            field="name",
            applicant_name_cn=cn,
            applicant_name_en=en,
            applicant_address_cn="中文址",
            applicant_address_en="EN Addr",
            side="cn",
        )
        == cn
    )
    assert (
        resolve_same_as_applicant(
            "同申请公司 Same as Applicant",
            field="name",
            applicant_name_cn=cn,
            applicant_name_en=en,
            applicant_address_cn="中文址",
            applicant_address_en="EN Addr",
            side="en",
        )
        == en
    )
    assert (
        resolve_same_as_applicant(
            "Same as Applicant",
            field="address",
            applicant_name_cn=cn,
            applicant_name_en=en,
            applicant_address_cn="中文址",
            applicant_address_en="",
            side="en",
        )
        == ""
    )


def test_zf_application_report_title_en_is_english():
    path = (
        Path("example")
        / "A2260437970101采埃孚年度 康"
        / "1.接样组"
        / "A22604379701--申请表更新.xlsx"
    )
    if not path.exists():
        return
    wb = load_workbook(path, data_only=True)
    sheet = wb[wb.sheetnames[0]]
    (
        _name_cn,
        name_en,
        _addr_cn,
        addr_en,
        _app_no,
        _rtn_cn,
        rtn_en,
        _rta_cn,
        rta_en,
    ) = parse_application_sheet1(sheet)
    assert "ZF" in name_en
    assert "Gangcheng" in addr_en or "Avenue" in addr_en
    assert rtn_en == name_en
    assert rta_en == addr_en


def test_bilingual_cover_customer_uses_en_not_han_title(tmp_path):
    """Report-title EN wrongly filled with Chinese must not overwrite applicant EN."""
    state = ProjectState(
        project_id="A2260999000101",
        applicant_name="中文委托方",
        applicant_name_en="EN Customer Co., Ltd.",
        applicant_address="中文地址",
        applicant_address_en="EN Address Line",
        report_title_name="中文委托方",
        report_title_address="中文地址",
        # Mis-resolved same-as (pre-fix shape): Han in EN slots
        report_title_name_en="中文委托方",
        report_title_address_en="中文地址",
        application_fields={"申请公司": "中文委托方", "申请公司地址": "中文地址"},
        application_fields_en={
            "申请公司": "EN Customer Co., Ltd.",
            "申请公司地址": "EN Address Line",
        },
    )
    template = Path("templates/template_ze.docx")
    assert template.exists()
    out = Path(tmp_path) / "ze_cover.docx"
    WordGenerator(str(template)).generate(
        state, str(out), report_language="中英文"
    )
    doc = Document(str(out))
    cover = doc.tables[0]
    assert cover.rows[0].cells[1].text.strip() == "中文委托方"
    assert cover.rows[1].cells[1].text.strip() == "EN Customer Co., Ltd."
    assert cover.rows[2].cells[1].text.strip() == "中文地址"
    assert cover.rows[3].cells[1].text.strip() == "EN Address Line"
    assert cover.rows[1].cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert cover.rows[3].cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
