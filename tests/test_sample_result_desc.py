import sys

from PySide6.QtCore import Qt
from PySide6.QtWidgets import QApplication, QLineEdit, QWidget

from src.generators.word_engine import WordGenerator
from src.models.project_state import (
    ProjectState,
    SampleStandardResult,
    TestLeg,
    TestNode,
    TestResult,
    TestSample,
    TestStandard,
)
from src.ui.test_detail_dialog import TestDetailDialog


def _app():
    app = QApplication.instance()
    if app is None:
        app = QApplication(sys.argv)
    return app


def _check_first_standard(dlg):
    chk = dlg.std_table.item(0, 0)
    assert chk is not None
    chk.setCheckState(Qt.Checked)


def _check_standards(dlg, *rows):
    for row in rows:
        chk = dlg.std_table.item(row, 0)
        assert chk is not None
        chk.setCheckState(Qt.Checked)


def test_sample_table_has_result_desc_column_and_autofills_on_generate():
    _app()
    desc = "试验后模块保持完整性。"
    standards = [
        {
            "标准号": "STD-1",
            "章节号": "1.1",
            "试验名称": "机械冲击试验",
            "标准描述": "条件",
            "结果描述": desc,
            "评价要求": "要求",
        }
    ]
    dlg = TestDetailDialog(TestNode(test_name="机械冲击"), standards, [])
    assert dlg.table.columnCount() == 4
    assert [dlg.table.horizontalHeaderItem(i).text() for i in range(4)] == [
        "",
        "样品编号",
        "结果描述",
        "测试结果",
    ]

    _check_first_standard(dlg)
    dlg.txt_sample_prefix.setText("A")
    dlg.txt_sample_start.setText("01")
    dlg.txt_sample_qty.setText("2")
    dlg._generate_samples()

    assert dlg.table.rowCount() == 2
    for row in range(2):
        desc_w = dlg.table.cellWidget(row, 2)
        assert isinstance(desc_w, QLineEdit)
        assert desc_w.text() == desc
        assert not desc_w.isReadOnly()


def test_per_sample_result_desc_editable_and_persisted():
    _app()
    desc = "标准结果描述"
    standards = [
        {
            "标准号": "STD-2b",
            "章节号": "2.3",
            "试验名称": "高温",
            "结果描述": desc,
        }
    ]
    dlg = TestDetailDialog(TestNode(test_name="高温"), standards, [])
    _check_first_standard(dlg)
    dlg.add_sample_row("A01", TestResult.PASS)
    dlg.add_sample_row("A02", TestResult.PASS)

    custom = "A01 单独改过的描述"
    dlg.table.cellWidget(0, 2).setText(custom)
    assert dlg.table.cellWidget(1, 2).text() == desc

    dlg.save_and_close()
    assert dlg.node_data.samples[0].result_desc == custom
    assert dlg.node_data.samples[1].result_desc == desc


def test_standard_result_desc_edit_still_overwrites_all_sample_rows():
    _app()
    desc = "标准结果描述"
    standards = [
        {
            "标准号": "STD-2c",
            "章节号": "2.4",
            "试验名称": "高温",
            "结果描述": desc,
        }
    ]
    dlg = TestDetailDialog(TestNode(test_name="高温"), standards, [])
    _check_first_standard(dlg)
    dlg.add_sample_row("A01", TestResult.PASS)
    dlg.add_sample_row("A02", TestResult.PASS)
    dlg.table.cellWidget(0, 2).setText("样品单独修改")

    updated = "上面改过的标准描述"
    editor = dlg.result_desc_table.cellWidget(0, 1)
    editor.setPlainText(updated)

    assert dlg.table.cellWidget(0, 2).text() == updated
    assert dlg.table.cellWidget(1, 2).text() == updated


def test_save_persists_sample_result_desc():
    _app()
    desc = "保存时应写入样品。"
    standards = [
        {
            "标准号": "STD-2",
            "章节号": "2.2",
            "试验名称": "高温",
            "结果描述": desc,
        }
    ]
    dlg = TestDetailDialog(TestNode(test_name="高温"), standards, [])
    _check_first_standard(dlg)
    dlg.add_sample_row("A01", TestResult.PASS)
    dlg.save_and_close()
    assert len(dlg.node_data.samples) == 1
    assert dlg.node_data.samples[0].result_desc == desc
    assert len(dlg.node_data.samples[0].standard_results) == 1
    assert dlg.node_data.samples[0].standard_results[0].result_desc == desc


def test_import_from_preceding_test_copies_ids_only():
    _app()
    desc_current = "本试验结果描述"
    desc_prev = "前置试验结果描述"
    standards = [
        {
            "标准号": "STD-3",
            "章节号": "3.3",
            "试验名称": "温湿度",
            "结果描述": desc_current,
        }
    ]
    prev_node = TestNode(
        test_name="高温",
        samples=[
            TestSample(sample_id="A01", result=TestResult.PASS, result_desc=desc_prev),
            TestSample(sample_id="A02", result=TestResult.FAIL, result_desc=desc_prev),
        ],
    )
    current_node = TestNode(test_name="温湿度", samples=[])
    state = ProjectState(
        project_id="P1",
        legs=[TestLeg(leg_id="L1", leg_name="Leg 1", nodes=[prev_node, current_node])],
    )

    class _Host(QWidget):
        def __init__(self):
            super().__init__()
            self.state = state

    host = _Host()
    dlg = TestDetailDialog(current_node, standards, [], host)
    _check_first_standard(dlg)
    assert dlg.btn_import_from_prev.isEnabled()

    dlg.add_sample_row("A01", TestResult.NA)
    dlg._import_from_preceding_test()

    assert dlg.table.rowCount() == 2
    id_widgets = [dlg.table.cellWidget(row, 1) for row in range(2)]
    assert [w.text() for w in id_widgets] == ["A01", "A02"]

    desc_widgets = [dlg.table.cellWidget(row, 2) for row in range(2)]
    assert all(w.text() == desc_current for w in desc_widgets)

    result_widgets = [dlg.table.cellWidget(row, 3) for row in range(2)]
    assert result_widgets[0].currentText() == TestResult.NA.value
    assert result_widgets[1].currentText() == TestResult.NA.value


def test_import_from_prev_disabled_for_first_node_in_leg():
    _app()
    node = TestNode(test_name="高温")
    state = ProjectState(
        project_id="P1",
        legs=[TestLeg(leg_id="L1", leg_name="Leg 1", nodes=[node])],
    )

    class _Host(QWidget):
        def __init__(self):
            super().__init__()
            self.state = state

    host = _Host()
    dlg = TestDetailDialog(node, [], [], host)
    assert not dlg.btn_import_from_prev.isEnabled()
    dlg.close()


def test_word_sample_table_includes_result_desc(tmp_path):
    from docx import Document

    template = tmp_path / "t.docx"
    Document().save(template)
    node = TestNode(
        test_name="冲击",
        result_desc="节点级描述",
        samples=[
            TestSample(
                sample_id="A01",
                result=TestResult.PASS,
                result_desc="样品级描述",
            )
        ],
    )
    gen = WordGenerator(str(template))
    doc = Document()
    anchor = doc.add_paragraph("ANCHOR")
    gen._insert_sample_result_table(doc, anchor, node)
    table = doc.tables[0]
    assert [c.text for c in table.rows[0].cells] == ["样品编号", "试验结果", "试验结论"]
    assert [c.text for c in table.rows[1].cells] == ["A01", "样品级描述", "合格"]


def test_word_multi_standard_emits_one_result_table_each(tmp_path):
    from docx import Document
    from docx.oxml.ns import qn

    template = tmp_path / "t.docx"
    Document().save(template)
    stds = [
        TestStandard(
            standard_id="S1",
            chapter="1",
            test_name="机械冲击试验",
            result_desc="冲击描述",
        ),
        TestStandard(
            standard_id="S2",
            chapter="2",
            test_name="湿热老化试验",
            result_desc="湿热描述",
        ),
    ]
    node = TestNode(
        test_name="组合",
        samples=[
            TestSample(
                sample_id="TP-262686912",
                result=TestResult.PASS,
                result_desc="拼接文",
                standard_results=[
                    SampleStandardResult(
                        standard_id="S1", chapter="1", result=TestResult.PASS
                    ),
                    SampleStandardResult(
                        standard_id="S2", chapter="2", result=TestResult.FAIL
                    ),
                ],
            ),
            TestSample(
                sample_id="TP-262686913",
                result=TestResult.PASS,
                result_desc="拼接文",
                standard_results=[
                    SampleStandardResult(
                        standard_id="S1", chapter="1", result=TestResult.PASS
                    ),
                    SampleStandardResult(
                        standard_id="S2", chapter="2", result=TestResult.PASS
                    ),
                ],
            ),
        ],
    )
    node.apply_standards(stds)
    gen = WordGenerator(str(template))
    doc = Document()
    anchor = doc.add_paragraph("ANCHOR")
    gen._insert_sample_result_table(doc, anchor, node)

    assert len(doc.tables) == 2
    headers = ["样品编号", "试验结果", "试验结论"]
    assert [c.text for c in doc.tables[0].rows[0].cells] == headers
    assert [c.text for c in doc.tables[1].rows[0].cells] == headers
    assert [c.text for c in doc.tables[0].rows[1].cells] == [
        "TP-262686912",
        "冲击描述",
        "合格",
    ]
    assert [c.text for c in doc.tables[0].rows[2].cells] == [
        "TP-262686913",
        "冲击描述",
        "合格",
    ]
    assert [c.text for c in doc.tables[1].rows[1].cells] == [
        "TP-262686912",
        "湿热描述",
        "不合格",
    ]
    assert [c.text for c in doc.tables[1].rows[2].cells] == [
        "TP-262686913",
        "湿热描述",
        "合格",
    ]

    # blank paragraph between the two tables
    t0 = doc.tables[0]._tbl
    t1 = doc.tables[1]._tbl
    between = []
    el = t0.getnext()
    while el is not None and el is not t1:
        between.append(el)
        el = el.getnext()
    assert between, "expected a blank paragraph between result tables"
    texts = []
    for el in between:
        texts.extend(t.text or "" for t in el.findall(".//" + qn("w:t")))
    assert "".join(texts).strip() == ""

    assert gen._node_conclusion(node) == "不合格"


def test_node_conclusion_aggregates_all_standard_results(tmp_path):
    from docx import Document

    template = tmp_path / "t.docx"
    Document().save(template)
    gen = WordGenerator(str(template))
    node = TestNode(
        test_name="组合",
        samples=[
            TestSample(
                sample_id="A01",
                result=TestResult.PASS,
                standard_results=[
                    SampleStandardResult(
                        standard_id="S1", chapter="1", result=TestResult.PASS
                    ),
                    SampleStandardResult(
                        standard_id="S2", chapter="2", result=TestResult.NA
                    ),
                ],
            )
        ],
    )
    assert gen._node_conclusion(node) == "合格"

    node.samples[0].standard_results[1].result = TestResult.FAIL
    assert gen._node_conclusion(node) == "不合格"


def _english_host(node):
    state = ProjectState(
        project_id="P1",
        edit_language="英文",
        legs=[TestLeg(leg_id="L1", leg_name="Leg 1", nodes=[node])],
    )

    class _Host(QWidget):
        def __init__(self):
            super().__init__()
            self.state = state

    return _Host()


def test_sample_result_desc_follows_edit_language():
    _app()
    desc_zh = "试验后模块保持完整性。"
    desc_en = "The module remained intact after the test."
    standards = [
        {
            "标准号": "STD-1",
            "章节号": "1.1",
            "试验名称": "机械冲击试验",
            "结果描述": desc_zh,
            "result": desc_en,
        }
    ]
    node = TestNode(
        test_name="机械冲击",
        samples=[TestSample(sample_id="A01", result=TestResult.PASS, result_desc=desc_zh)],
        standards=[
            TestStandard(
                standard_id="STD-1",
                chapter="1.1",
                test_name="机械冲击试验",
                result_desc=desc_zh,
                result_desc_en=desc_en,
            )
        ],
    )
    host = _english_host(node)
    dlg = TestDetailDialog(node, standards, [], host)
    desc_w = dlg.table.cellWidget(0, 2)
    assert desc_w is not None
    assert desc_w.text() == desc_en
    dlg.close()


def test_bulk_result_combo_includes_na_without_scroll():
    _app()
    dlg = TestDetailDialog(TestNode(test_name="高温"), [], [])
    texts = [dlg.combo_bulk_result.itemText(i) for i in range(dlg.combo_bulk_result.count())]
    assert texts == ["—", "合格", "不合格", "N/A"]
    assert dlg.combo_bulk_result.maxVisibleItems() >= 4
    dlg.close()


def test_multi_standard_splits_sample_tables_without_joining_desc():
    _app()
    desc_a = "冲击描述AAA"
    desc_b = "湿热描述BBB"
    standards = [
        {
            "标准号": "S1",
            "章节号": "1",
            "试验名称": "机械冲击试验",
            "结果描述": desc_a,
        },
        {
            "标准号": "S2",
            "章节号": "2",
            "试验名称": "湿热老化试验",
            "结果描述": desc_b,
        },
    ]
    dlg = TestDetailDialog(TestNode(test_name="组合"), standards, [])
    _check_standards(dlg, 0, 1)
    assert len(dlg._sample_table_slots) == 2
    dlg.add_sample_row("A01", TestResult.PASS)

    t0 = dlg._sample_table_slots[0]["table"]
    t1 = dlg._sample_table_slots[1]["table"]
    assert t0.rowCount() == 1 and t1.rowCount() == 1
    assert t0.cellWidget(0, 2).text() == desc_a
    assert t1.cellWidget(0, 2).text() == desc_b
    assert desc_b not in t0.cellWidget(0, 2).text()
    assert t0.cellWidget(0, 1).text() == "A01"
    assert t1.cellWidget(0, 1).text() == "A01"


def test_multi_standard_independent_conclusions_save_and_reload():
    _app()
    standards = [
        {
            "标准号": "S1",
            "章节号": "1",
            "试验名称": "机械冲击试验",
            "结果描述": "冲击描述",
        },
        {
            "标准号": "S2",
            "章节号": "2",
            "试验名称": "湿热老化试验",
            "结果描述": "湿热描述",
        },
    ]
    dlg = TestDetailDialog(TestNode(test_name="组合"), standards, [])
    _check_standards(dlg, 0, 1)
    dlg.add_sample_row("A01", TestResult.NA)

    t0 = dlg._sample_table_slots[0]["table"]
    t1 = dlg._sample_table_slots[1]["table"]
    t0.cellWidget(0, 3).setCurrentIndex(t0.cellWidget(0, 3).findData(TestResult.PASS))
    t1.cellWidget(0, 3).setCurrentIndex(t1.cellWidget(0, 3).findData(TestResult.FAIL))
    dlg.save_and_close()

    sample = dlg.node_data.samples[0]
    assert sample.sample_id == "A01"
    assert len(sample.standard_results) == 2
    assert sample.standard_results[0].result == TestResult.PASS
    assert sample.standard_results[1].result == TestResult.FAIL
    assert sample.result == TestResult.PASS

    dlg2 = TestDetailDialog(dlg.node_data, standards, [])
    assert len(dlg2._sample_table_slots) == 2
    assert dlg2._sample_table_slots[0]["table"].cellWidget(0, 3).currentData() == TestResult.PASS
    assert dlg2._sample_table_slots[1]["table"].cellWidget(0, 3).currentData() == TestResult.FAIL
    dlg2.close()


def test_legacy_scalar_result_fills_all_tables_on_load():
    _app()
    standards = [
        {
            "标准号": "S1",
            "章节号": "1",
            "试验名称": "机械冲击试验",
            "结果描述": "冲击描述",
        },
        {
            "标准号": "S2",
            "章节号": "2",
            "试验名称": "湿热老化试验",
            "结果描述": "湿热描述",
        },
    ]
    node = TestNode(
        test_name="组合",
        samples=[TestSample(sample_id="A01", result=TestResult.FAIL, result_desc="旧拼接")],
        standards=[
            TestStandard(
                standard_id="S1", chapter="1", test_name="机械冲击试验", result_desc="冲击描述"
            ),
            TestStandard(
                standard_id="S2", chapter="2", test_name="湿热老化试验", result_desc="湿热描述"
            ),
        ],
    )
    dlg = TestDetailDialog(node, standards, [])
    assert len(dlg._sample_table_slots) == 2
    assert dlg._sample_table_slots[0]["table"].cellWidget(0, 3).currentData() == TestResult.FAIL
    assert dlg._sample_table_slots[1]["table"].cellWidget(0, 3).currentData() == TestResult.FAIL
    assert dlg._sample_table_slots[0]["table"].cellWidget(0, 2).text() == "冲击描述"
    assert dlg._sample_table_slots[1]["table"].cellWidget(0, 2).text() == "湿热描述"
    dlg.close()
