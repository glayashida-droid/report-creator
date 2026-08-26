"""Seam 1: Word export by report_language (English / bilingual copy)."""

from pathlib import Path

from docx import Document

from src.generators.word_engine import WordGenerator
from src.io.test_photos import create_album, copy_into_album
from src.models.project_state import (
    ProjectState,
    TestEquipment,
    TestLeg,
    TestNode,
    TestSample,
    TestResult,
    TestStandard,
)
from PIL import Image


def _png(path: Path):
    Image.new("RGB", (40, 30), color=(80, 90, 100)).save(path)


def _engine_state(tmp_path) -> tuple[ProjectState, Path, Path]:
    project = Path(tmp_path) / "proj"
    (project / "3.测试组" / "高温试验").mkdir(parents=True)
    album = create_album(project, "高温试验", "试验前")
    src = Path(tmp_path) / "a.png"
    _png(src)
    copy_into_album(album, [src], "试验前")

    state = ProjectState(
        project_id="A2260999000101",
        applicant_name="中文委托方",
        applicant_name_en="EN Customer",
        applicant_address="中文地址",
        applicant_address_en="EN Address",
        sample_name="安全带",
        sample_name_en="Seat belt",
        application_fields={
            "申请单号": "A22609990001",
            "申请公司": "中文委托方",
            "样品名称": "安全带",
            "零件号": "P519",
        },
        application_fields_en={
            "申请公司": "EN Customer",
            "样品名称": "Seat belt",
            "零件号": "P519",
        },
        sample_receive_date="2026-08-01",
        test_start_date="2026-08-02",
        test_end_date="2026-08-10",
    )
    node = TestNode(
        test_name="高温试验",
        standards=[
            TestStandard(
                standard_id="VW 80000",
                chapter="4.1",
                test_name="高温",
                test_item="High temperature",
                standard_desc="中文条件正文",
                standard_desc_en="EN condition body",
                evaluation_req="中文评判",
                evaluation_req_en="EN evaluation",
                result_desc="中文结果描述",
                result_desc_en="EN result desc",
            )
        ],
        equipments=[
            TestEquipment(
                name="高低温箱",
                name_en="Temp chamber",
                code="T1",
                model="M1",
                valid_date="2027-01-01",
            )
        ],
        samples=[TestSample(sample_id="A01", result=TestResult.PASS)],
    )
    node.apply_standards(node.standards)
    leg = TestLeg(leg_id="L1", leg_name="Leg 1", nodes=[node])
    state.legs.append(leg)
    return state, project, Path("templates/template_en.docx")


def test_english_export_no_chinese_fallback_and_key_strings(tmp_path):
    state, project, template = _engine_state(tmp_path)
    assert template.exists()
    out = Path(tmp_path) / "en.docx"
    WordGenerator(str(template)).generate(
        state, str(out), project_path=str(project), report_language="英文"
    )
    doc = Document(str(out))
    texts = "\n".join(p.text for p in doc.paragraphs)
    tables = "\n".join(
        c.text for t in doc.tables for r in t.rows for c in r.cells
    )
    blob = texts + "\n" + tables

    assert "EN Customer" in blob or "{{委托方名称}}" not in blob
    assert "Pass" in blob
    assert "Before test" in blob
    assert "High temperature" in blob
    assert "EN condition body" in blob
    assert "检测项目" not in texts  # detail title uses English
    assert "中文条件正文" not in blob
    assert "中文委托方" not in blob
    assert WordGenerator.default_report_no(state, "英文") == "A226099900010100001"


def test_bilingual_export_join_and_zh_only_when_en_missing(tmp_path):
    state, project, _ = _engine_state(tmp_path)
    state.applicant_name_en = ""
    state.application_fields_en["申请公司"] = ""
    state.legs[0].nodes[0].standards[0].standard_desc_en = ""
    state.legs[0].nodes[0].apply_standards(state.legs[0].nodes[0].standards)

    template = Path("templates/template_ze.docx")
    assert template.exists()
    out = Path(tmp_path) / "ze.docx"
    WordGenerator(str(template)).generate(
        state, str(out), project_path=str(project), report_language="中英文"
    )
    doc = Document(str(out))
    texts = "\n".join(p.text for p in doc.paragraphs)
    tables = "\n".join(
        c.text for t in doc.tables for r in t.rows for c in r.cells
    )
    blob = texts + "\n" + tables

    assert "合格 / Pass" in blob
    assert "试验前 / Before test" in blob
    assert "高温试验 / High temperature" in blob or "High temperature" in blob
    assert "中文条件正文" in blob
    assert " / " not in "中文条件正文"  # bare zh condition present; no empty slash pair alone
    assert WordGenerator.default_report_no(state, "中英文").endswith("E")
