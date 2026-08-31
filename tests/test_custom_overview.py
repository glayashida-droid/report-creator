"""Manual project overview rows: editable label/value and report export."""

from pathlib import Path

from docx import Document

from src.generators.word_engine import WordGenerator
from src.models.project_state import CustomOverviewField, ProjectState


def test_custom_overview_fields_persist_and_export(tmp_path):
    state = ProjectState(
        project_id="A1",
        application_fields={"样品名称": "测试样品"},
        custom_overview_fields=[
            CustomOverviewField(
                id="abc123",
                label_cn="备注",
                value_cn="客户补充说明",
            )
        ],
    )

    key = ProjectState.custom_overview_key("abc123")
    rows = dict(state.iter_overview_fields("中文"))
    assert rows[key] == "客户补充说明"
    assert state.overview_display_label(key, "中文") == "备注"

    json_path = tmp_path / "state.json"
    state.save_to_file(str(json_path))
    loaded = ProjectState.load_from_file(str(json_path))
    assert len(loaded.custom_overview_fields) == 1
    assert loaded.custom_overview_fields[0].value_cn == "客户补充说明"

    template = Path("templates/template_ze.docx")
    assert template.is_file(), "run tests from project root"
    out = tmp_path / "out.docx"
    WordGenerator(str(template)).generate(state, str(out), report_language="中文")

    doc = Document(str(out))
    all_cells = [
        (row.cells[0].text or "").strip()
        for table in doc.tables
        for row in table.rows
        if len(row.cells) >= 2
    ]
    all_values = [
        (row.cells[1].text or "").strip()
        for table in doc.tables
        for row in table.rows
        if len(row.cells) >= 2
    ]
    assert "备注" in all_cells
    assert any("客户补充说明" in value for value in all_values)


def test_custom_overview_bilingual_sides(tmp_path):
    state = ProjectState(project_id="A1")
    row = state.add_custom_overview_field()
    state.set_custom_overview_label(row.id, "批次号")
    state.set_custom_overview_value(row.id, "B-001")
    state.edit_language = "英文"
    state.set_custom_overview_label(row.id, "Batch No.")
    state.set_custom_overview_value(row.id, "B-001")

    key = ProjectState.custom_overview_key(row.id)
    assert dict(state.iter_overview_fields("中文"))[key] == "B-001"
    assert dict(state.iter_overview_fields("英文"))[key] == "B-001"
    assert state.overview_display_label(key, "英文") == "Batch No."


def test_remove_custom_overview_field():
    state = ProjectState(project_id="A1")
    row = state.add_custom_overview_field()
    state.set_custom_overview_label(row.id, "临时")
    state.set_custom_overview_value(row.id, "值")
    key = ProjectState.custom_overview_key(row.id)
    assert key in dict(state.iter_overview_fields("中文"))

    state.remove_custom_overview_field(row.id)
    assert not state.custom_overview_fields
    assert key not in dict(state.iter_overview_fields("中文"))
