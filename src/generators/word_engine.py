"""Render ProjectState into a Word report from templates/template_zh.docx."""

from __future__ import annotations

import io
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips
from docx.text.paragraph import Paragraph
from openpyxl.utils.cell import range_boundaries

from src.io.data_tables import (
    infer_header_row_count,
    list_attachment_refs,
    read_preview_snapshot,
    resolve_attachment_path,
)
from src.io.test_photos import list_albums, list_photos, uses_data_photo_layout
from src.language_copy import (
    field_label,
    format_conclusion,
    has_chinese,
    language_text,
    photo_caption,
    table_header_label,
)
from src.models.project_state import (
    DataTableRef,
    ProjectState,
    TestLeg,
    TestNode,
    TestResult,
    TestSample,
)

# Fixed section-(1) text kept in the Chinese template contract
ENV_CONDITION_TEXT = "（23±5）℃，（50±25）%RH"

PHOTO_WIDTH_IN = 2.95
DATA_PHOTO_WIDTH_IN = 5.5
CONDITION_IMAGE_MAX_WIDTH_IN = 5.0
# Avoid stretching tiny Excel-cell PNGs past ~150 DPI (reduces blur)
CONDITION_IMAGE_MIN_DPI = 150
# Match Word「默认解析度」330 ppi; q=96 ≈ golden photo media, slightly above
EMBED_PPI = 330
EMBED_JPEG_QUALITY = 96

# Report typography (matches CTI golden reports)
FONT_EAST = "方正书宋简体"
FONT_LATIN = "Times New Roman"
SIZE_BODY = 10.5  # 五号
SIZE_CAPTION = 9  # 小五
SIZE_PHOTO_TITLE = 18
# Golden body: line=276 (≈1.15) + snapToGrid off. Template has docGrid
# linePitch=312; without snapToGrid=0 Word snaps to the grid and looks double-spaced.
LINE_SPACING_TWIPS = 276
# Cover key-value row height (golden ≈ 170; keep readable for 五号)
COVER_ROW_HEIGHT_TWIPS = 200

# Usable page width in DXA (≈ golden cover table total)
_CONTENT_WIDTH_DXA = 8802

# Column widths from golden A226061368610100001E.docx (DXA)
_WIDTHS_COVER_INFO = (3423, 5292)  # ~39:61, not 50:50
_WIDTHS_SAMPLE_LIST = (704, 1697, 1276, 1236, 1980)
_WIDTHS_SUMMARY = (704, 2090, 1690, 1400, 1500, 1418)
_WIDTHS_EQUIPMENT = (704, 2693, 1843, 2126, 1418)
_WIDTHS_SAMPLE_RESULT = (1523, 6269, 1353)

# Homepage overview keys that belong in the dynamic 样品信息表 (label as-is).
# Dates / quantity / see-below rows are appended by the engine.
_COVER_SKIP = {
    "申请单号",
    "申请公司",
    "申请公司地址",
    "报告抬头公司",
    "报告抬头地址",
    "送样数量",
}


class WordGenerator:
    def __init__(self, template_path: str):
        self.template_path = template_path
        self._report_language = "中文"

    def generate(
        self,
        state: ProjectState,
        output_path: str,
        project_path: str = None,
        leg_filter: str = None,
        report_language: str = "中文",
        report_no: str | None = None,
    ):
        self._report_language = (report_language or "中文").strip() or "中文"
        doc = Document(self.template_path)
        placeholders = self._build_placeholders(
            state, report_language=self._report_language, report_no=report_no
        )
        self._replace_everywhere(doc, placeholders)
        self._tighten_cover_info_table(doc)
        if self._report_language == "中英文":
            self._fill_ze_cover_english(doc, state)

        nodes = list(state.iter_nodes_for_export(leg_filter))

        self._replace_marker_with_table(
            doc, "{{样品信息表}}", lambda d, p: self._insert_sample_info_table(d, p, state)
        )
        self._replace_marker_with_table(
            doc, "{{样品清单表}}", lambda d, p: self._insert_sample_list_table(d, p, state, [n for _, n in nodes])
        )
        self._replace_marker_with_table(
            doc, "{{结果汇总表}}", lambda d, p: self._insert_summary_table(d, p, [n for _, n in nodes])
        )
        self._replace_marker_with_blocks(
            doc,
            "{{试验明细}}",
            lambda d, p: self._insert_test_details(d, p, nodes, project_path),
        )

        doc.save(output_path)

    def _lang(self) -> str:
        return self._report_language or "中文"

    def _side_lang(self) -> str:
        """Language side for overview field values (中英文 uses Chinese rows + separate EN fill)."""
        lang = self._lang()
        if lang == "英文":
            return "英文"
        return "中文"

    def _L(self, zh: str, en: str) -> str:
        """Inline section labels (raw concat for 中英文)."""
        return table_header_label(zh, en, self._lang(), inline=True)

    def _H(self, zh: str, en: str) -> str:
        """Word table column headers (bilingual line break)."""
        return table_header_label(zh, en, self._lang())

    def _or_slash(self, text: str) -> str:
        """Chinese/bilingual keep '/' placeholder; English stays empty when missing."""
        text = (text or "").strip()
        if text:
            return text
        return "" if self._lang() == "英文" else "/"

    def _fill_ze_cover_english(self, doc: Document, state: ProjectState):
        """Write Customer / Address value cells on template_ze cover (no placeholders)."""
        en_name = (state.applicant_name_en or "").strip() or (
            (state.application_fields_en or {}).get("申请公司") or ""
        ).strip()
        en_addr = (state.applicant_address_en or "").strip() or (
            (state.application_fields_en or {}).get("申请公司地址") or ""
        ).strip()
        title_en = (state.report_title_name_en or "").strip()
        title_addr_en = (state.report_title_address_en or "").strip()
        # Prefer report-title EN when present and actually English (no Han)
        if title_en and not has_chinese(title_en):
            en_name = title_en
        if title_addr_en and not has_chinese(title_addr_en):
            en_addr = title_addr_en
        if has_chinese(en_name):
            en_name = ""
        if has_chinese(en_addr):
            en_addr = ""
        for table in doc.tables:
            if len(table.rows) < 4 or len(table.columns) < 2:
                continue
            c0 = (table.rows[1].cells[0].text or "").strip()
            c2 = (table.rows[3].cells[0].text or "").strip()
            if "Customer" in c0 and "Address" in c2:
                self._set_cell_text(
                    table.rows[1].cells[1], en_name, align=WD_ALIGN_PARAGRAPH.LEFT
                )
                self._set_cell_text(
                    table.rows[3].cells[1], en_addr, align=WD_ALIGN_PARAGRAPH.LEFT
                )
                break

    # ------------------------------------------------------------------ placeholders

    def _build_placeholders(
        self,
        state: ProjectState,
        report_language: str = "中文",
        report_no: str | None = None,
    ) -> Dict[str, str]:
        lang = (report_language or "中文").strip() or "中文"
        # 中英文 cover: CN placeholders for 委托单位/地址; EN filled separately
        side = "英文" if lang == "英文" else "中文"
        visible = dict(state.iter_overview_fields(side))
        if report_no is None:
            report_no = self.default_report_no(state, report_language)
        else:
            report_no = (report_no or "").strip()
        name_fallback = state.applicant_name_en if side == "英文" else state.applicant_name
        addr_fallback = state.applicant_address_en if side == "英文" else state.applicant_address
        sample_fallback = state.sample_name_en if side == "英文" else state.sample_name
        # Report-title preference for cover name/address when set
        if side == "中文":
            title_name = (state.report_title_name or "").strip()
            title_addr = (state.report_title_address or "").strip()
            if title_name:
                name_fallback = title_name
            if title_addr:
                addr_fallback = title_addr
        else:
            title_name = (state.report_title_name_en or "").strip()
            title_addr = (state.report_title_address_en or "").strip()
            if title_name:
                name_fallback = title_name
            if title_addr:
                addr_fallback = title_addr
        placeholders = {
            "{{委托方名称}}": visible.get("申请公司", "") or name_fallback or "",
            "{{委托方地址}}": visible.get("申请公司地址", "") or addr_fallback or "",
            "{{样品名称}}": visible.get("样品名称", "") or sample_fallback or "",
            "{{样品接收日期}}": self._fmt_date(state.sample_receive_date),
            "{{检测开始日期}}": self._fmt_date(state.test_start_date),
            "{{检测结束日期}}": self._fmt_date(state.test_end_date),
            "{{申请单号}}": visible.get("申请单号", ""),
            "{{报告抬头公司}}": visible.get("报告抬头公司", ""),
            "{{报告抬头地址}}": visible.get("报告抬头地址", ""),
            "{{报告编号}}": report_no,
        }
        for key, val in visible.items():
            ph_key = (
                state.overview_display_label(key, side)
                if ProjectState.is_custom_overview_key(key)
                else key
            )
            if ph_key:
                placeholders["{{%s}}" % ph_key] = val
        for key in state.excluded_overview_keys or []:
            placeholders.setdefault("{{%s}}" % key, "")
            if key == "申请公司":
                placeholders["{{委托方名称}}"] = ""
            elif key == "申请公司地址":
                placeholders["{{委托方地址}}"] = ""
            elif key == "样品名称":
                placeholders["{{样品名称}}"] = ""
        return placeholders

    @staticmethod
    def default_report_no(state: ProjectState, report_language: str = "中文") -> str:
        pid = (state.project_id or "").strip()
        if not pid:
            app = (state.application_fields or {}).get("申请单号") or ""
            pid = str(app).strip()
        if not pid:
            return ""
        base = f"{pid}00001"
        lang = (report_language or "中文").strip()
        if lang == "英文":
            return base
        if lang == "中英文":
            return f"{base}E"
        return f"{base}C"  # 中文

    @staticmethod
    def report_filename_stem(report_no: str) -> str:
        from src.io.data_tables import sanitize_filename_stem

        stem = sanitize_filename_stem(report_no)
        if stem == "未命名数据表":
            return "未命名报告"
        return stem

    @staticmethod
    def next_duplicate_report_path(folder: Path, stem: str) -> Path:
        """First available stem-2.docx, stem-3.docx, … under folder."""
        n = 2
        while True:
            candidate = folder / f"{stem}-{n}.docx"
            if not candidate.exists():
                return candidate
            n += 1

    @staticmethod
    def _fmt_date(value: Optional[str]) -> str:
        text = (value or "").strip()
        if not text:
            return ""
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
            try:
                return datetime.strptime(text[:10], fmt).strftime("%Y.%m.%d")
            except ValueError:
                continue
        return text.replace("-", ".").replace("/", ".")

    @staticmethod
    def _fmt_period(start: Optional[str], end: Optional[str]) -> str:
        a = WordGenerator._fmt_date(start)
        b = WordGenerator._fmt_date(end)
        if a and b:
            return f"{a}~{b}"
        return a or b or ""

    # ------------------------------------------------------------------ replace helpers

    def _replace_everywhere(self, doc: Document, placeholders: Dict[str, str]):
        for p in doc.paragraphs:
            self._replace_text_in_paragraph(p, placeholders)
        for table in doc.tables:
            self._replace_in_table(table, placeholders)
        for section in doc.sections:
            for part in (section.header, section.footer):
                for p in part.paragraphs:
                    self._replace_text_in_paragraph(p, placeholders)
                for table in part.tables:
                    self._replace_in_table(table, placeholders)

    def _replace_in_table(self, table, placeholders: Dict[str, str]):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    self._replace_text_in_paragraph(p, placeholders)
                for nested in cell.tables:
                    self._replace_in_table(nested, placeholders)

    def _replace_text_in_paragraph(self, paragraph: Paragraph, placeholders: Dict[str, str]):
        if not placeholders:
            return
        # Never smash paragraphs that contain PAGE/NUMPAGES fields
        if paragraph._element.findall(".//" + qn("w:fldChar")):
            self._replace_in_runs_only(paragraph, placeholders)
            return

        text = paragraph.text
        if not text:
            return
        new_text = text
        for key, val in placeholders.items():
            if key in new_text:
                new_text = new_text.replace(key, str(val))
        if new_text == text:
            return
        if self._replace_in_runs_only(paragraph, placeholders):
            return
        # Fallback: rewrite first run, clear the rest (keeps basic style)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = new_text

    @staticmethod
    def _replace_in_runs_only(paragraph: Paragraph, placeholders: Dict[str, str]) -> bool:
        """Replace placeholders that sit entirely inside a single run. Returns True if any changed."""
        changed = False
        for run in paragraph.runs:
            text = run.text or ""
            new_text = text
            for key, val in placeholders.items():
                if key in new_text:
                    new_text = new_text.replace(key, str(val))
            if new_text != text:
                run.text = new_text
                changed = True
        return changed

    def _find_paragraph(self, doc: Document, marker: str) -> Optional[Paragraph]:
        for p in doc.paragraphs:
            if marker in (p.text or ""):
                return p
        return None

    def _replace_marker_with_table(self, doc: Document, marker: str, builder):
        p = self._find_paragraph(doc, marker)
        if p is None:
            return
        builder(doc, p)
        self._delete_paragraph(p)

    def _replace_marker_with_blocks(self, doc: Document, marker: str, builder):
        p = self._find_paragraph(doc, marker)
        if p is None:
            return
        builder(doc, p)
        self._delete_paragraph(p)

    @staticmethod
    def _delete_paragraph(paragraph: Paragraph):
        el = paragraph._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    def _insert_before(self, anchor: Paragraph, element):
        anchor._element.addprevious(element)

    def _add_para_before(
        self,
        doc: Document,
        anchor: Paragraph,
        text: str = "",
        *,
        bold: bool = False,
        size: float = SIZE_BODY,
        align=None,
    ) -> Paragraph:
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        if text:
            run = p.add_run(text)
            self._style_run(run, bold=bold, size=size)
        self._apply_tight_spacing(p)
        self._insert_before(anchor, p._element)
        # python-docx appends to body; we already moved the element
        return p

    def _add_page_break_before(self, doc: Document, anchor: Paragraph) -> Paragraph:
        p = self._add_para_before(doc, anchor, "")
        p.add_run().add_break(WD_BREAK.PAGE)
        return p

    def _add_table_before(
        self,
        doc: Document,
        anchor: Paragraph,
        rows: int,
        cols: int,
        *,
        bordered: bool = True,
    ):
        table = doc.add_table(rows=rows, cols=cols)
        if bordered:
            try:
                table.style = "Table Grid"
            except KeyError:
                pass
        else:
            try:
                table.style = "Normal Table"
            except KeyError:
                pass
            self._set_table_borders(table, None)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            self._set_row_cant_split(row)
        self._insert_before(anchor, table._tbl)
        return table

    @staticmethod
    def _apply_tight_spacing(paragraph: Paragraph) -> None:
        """Match golden body: 0 before/after, ~1.15 line spacing, no doc-grid snap."""
        pf = paragraph.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = LINE_SPACING_TWIPS / 240.0
        pPr = paragraph._element.get_or_add_pPr()

        def _flag(tag: str, val: str = "0") -> None:
            el = pPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                pPr.append(el)
            el.set(qn("w:val"), val)

        _flag("w:snapToGrid")
        _flag("w:adjustRightInd")
        _flag("w:autoSpaceDE")
        _flag("w:autoSpaceDN")

        sp = pPr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            pPr.append(sp)
        sp.set(qn("w:before"), "0")
        sp.set(qn("w:after"), "0")
        sp.set(qn("w:beforeAutospacing"), "0")
        sp.set(qn("w:afterAutospacing"), "0")
        sp.set(qn("w:line"), str(LINE_SPACING_TWIPS))
        sp.set(qn("w:lineRule"), "auto")

    @staticmethod
    def _style_run(
        run,
        *,
        bold: bool = False,
        size: float = SIZE_BODY,
        east: str = FONT_EAST,
    ):
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = FONT_LATIN
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), FONT_LATIN)
        rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        rFonts.set(qn("w:eastAsia"), east)

    def _set_cell_text(
        self,
        cell,
        text: str,
        *,
        bold: bool = False,
        size: float = SIZE_BODY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    ):
        raw = "" if text is None else str(text)
        # Drop blank lines; keep intentional single newlines as soft breaks in one para
        lines = [ln.strip() for ln in raw.replace("\r\n", "\n").split("\n")]
        lines = [ln for ln in lines if ln]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # cell.text="" leaves one empty paragraph; remove extras left by prior merges
        for extra in list(cell.paragraphs)[1:]:
            el = extra._element
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
        p = cell.paragraphs[0]
        p.alignment = align
        self._apply_tight_spacing(p)
        # clear runs
        for run in list(p.runs):
            run._element.getparent().remove(run._element)
        if not lines:
            run = p.add_run("")
            self._style_run(run, bold=bold, size=size)
            return
        for i, line in enumerate(lines):
            if i:
                p.add_run().add_break()
            run = p.add_run(line)
            self._style_run(run, bold=bold, size=size)

    @staticmethod
    def _set_row_cant_split(row) -> None:
        """Keep the whole row on one page (Word: 允许跨页断行 = off)."""
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))

    @staticmethod
    def _set_row_as_tbl_header(row) -> None:
        """Mark row as repeating header (Word: 在各页顶端以标题行形式重复出现)."""
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        if trPr.find(qn("w:tblHeader")) is None:
            trPr.append(OxmlElement("w:tblHeader"))

    @staticmethod
    def _merged_slave_cells(snap) -> set:
        """(r,c) cells covered by a merge but not the top-left anchor."""
        covered = set()
        origin_r = snap.origin_row or 1
        origin_c = snap.origin_col or 1
        for merge in snap.merges or []:
            try:
                min_c, min_r, max_c, max_r = range_boundaries(merge)
            except Exception:
                continue
            r0 = min_r - origin_r
            c0 = min_c - origin_c
            r1 = max_r - origin_r
            c1 = max_c - origin_c
            for r in range(r0, r1 + 1):
                for c in range(c0, c1 + 1):
                    if (r, c) != (r0, c0):
                        covered.add((r, c))
        return covered

    @staticmethod
    def _apply_snapshot_merges(table, snap) -> None:
        """Replay Excel merges from PreviewSnapshot onto a Word table."""
        rows = len(snap.values or [])
        cols = max((len(r) for r in (snap.values or [])), default=0)
        if rows == 0 or cols == 0:
            return
        origin_r = snap.origin_row or 1
        origin_c = snap.origin_col or 1
        # Larger merges first so nested ranges are less likely to fail
        ranges = []
        for merge in snap.merges or []:
            try:
                min_c, min_r, max_c, max_r = range_boundaries(merge)
            except Exception:
                continue
            r0 = min_r - origin_r
            c0 = min_c - origin_c
            r1 = max_r - origin_r
            c1 = max_c - origin_c
            if r0 < 0 or c0 < 0 or r1 >= rows or c1 >= cols:
                continue
            if r1 > r0 or c1 > c0:
                ranges.append((r0, c0, r1, c1))
        ranges.sort(key=lambda t: -((t[2] - t[0] + 1) * (t[3] - t[1] + 1)))
        for r0, c0, r1, c1 in ranges:
            try:
                table.cell(r0, c0).merge(table.cell(r1, c1))
            except Exception:
                continue

    @staticmethod
    def _set_row_height(row, twips: int, *, exact: bool = False) -> None:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = trPr.find(qn("w:trHeight"))
        if trHeight is None:
            trHeight = OxmlElement("w:trHeight")
            trPr.append(trHeight)
        trHeight.set(qn("w:val"), str(int(twips)))
        trHeight.set(qn("w:hRule"), "exact" if exact else "atLeast")

    @staticmethod
    def _set_table_cell_margins(
        table, *, top: int = 0, left: int = 40, bottom: int = 0, right: int = 40
    ) -> None:
        """Cell margins in DXA (twips)."""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        existing = tblPr.find(qn("w:tblCellMar"))
        if existing is not None:
            tblPr.remove(existing)
        mar = OxmlElement("w:tblCellMar")
        for edge, val in (
            ("top", top),
            ("left", left),
            ("bottom", bottom),
            ("right", right),
        ):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:w"), str(int(val)))
            el.set(qn("w:type"), "dxa")
            mar.append(el)
        tblPr.append(mar)

    @staticmethod
    def _condition_image_width_in(blob: bytes) -> float:
        """Cap display width so small Excel-cell PNGs stay near CONDITION_IMAGE_MIN_DPI."""
        from PIL import Image

        try:
            with Image.open(io.BytesIO(blob)) as im:
                px_w = im.size[0]
        except Exception:
            return CONDITION_IMAGE_MAX_WIDTH_IN
        if px_w <= 0:
            return CONDITION_IMAGE_MAX_WIDTH_IN
        native_in = px_w / float(CONDITION_IMAGE_MIN_DPI)
        return max(0.8, min(CONDITION_IMAGE_MAX_WIDTH_IN, native_in))

    @staticmethod
    def _target_embed_px(width_in: float) -> int:
        return max(1, int(round(float(width_in) * EMBED_PPI)))

    @staticmethod
    def _prepare_embed_stream(source: "Path | bytes", width_in: float) -> io.BytesIO:
        """Downscale test photos to width_in×EMBED_PPI (never upscale), JPEG encode.

        Standard-library condition images are embedded elsewhere without this path.
        """
        from PIL import Image

        if isinstance(source, (bytes, bytearray)):
            raw = bytes(source)
        else:
            raw = Path(source).read_bytes()

        target_w = WordGenerator._target_embed_px(width_in)
        with Image.open(io.BytesIO(raw)) as im:
            im.load()
            src_format = (im.format or "").upper()
            w, h = im.size
            if w <= 0 or h <= 0:
                return io.BytesIO(raw)

            # Already at/under target: keep original JPEG bytes (no re-encode)
            if w <= target_w and src_format in {"JPEG", "JPG", "MPO"}:
                return io.BytesIO(raw)

            if w > target_w:
                new_h = max(1, int(round(h * (target_w / float(w)))))
                im = im.resize((target_w, new_h), Image.Resampling.LANCZOS)

            if im.mode in ("RGBA", "LA") or (
                im.mode == "P" and "transparency" in im.info
            ):
                rgba = im.convert("RGBA")
                bg = Image.new("RGB", rgba.size, (255, 255, 255))
                bg.paste(rgba, mask=rgba.split()[-1])
                im = bg
            elif im.mode != "RGB":
                im = im.convert("RGB")

            out = io.BytesIO()
            im.save(
                out,
                format="JPEG",
                quality=EMBED_JPEG_QUALITY,
                optimize=True,
                progressive=True,
            )
            out.seek(0)
            return out

    @staticmethod
    def _set_table_borders(table, border_val: Optional[str] = "single"):
        """Set all table borders. border_val=None clears borders (nil)."""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            if border_val is None:
                el.set(qn("w:val"), "nil")
            else:
                el.set(qn("w:val"), border_val)
                el.set(qn("w:sz"), "4")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "000000")
            borders.append(el)
        tblPr.append(borders)

    @staticmethod
    def _set_col_widths(table, widths_dxa: Sequence[int]):
        """Set fixed column widths (DXA / twips)."""
        cols = list(widths_dxa)
        if not cols:
            return
        tbl = table._tbl
        grid = tbl.find(qn("w:tblGrid"))
        if grid is None:
            grid = OxmlElement("w:tblGrid")
            tblPr = tbl.tblPr
            if tblPr is not None:
                tblPr.addnext(grid)
            else:
                tbl.insert(0, grid)
        for child in list(grid):
            grid.remove(child)
        for w in cols:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(int(w)))
            grid.append(gc)
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx >= len(cols):
                    break
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.append(tcW)
                tcW.set(qn("w:w"), str(int(cols[idx])))
                tcW.set(qn("w:type"), "dxa")
                cell.width = Twips(int(cols[idx]))

    @staticmethod
    def _estimate_col_widths(
        values: Sequence[Sequence[str]], total_dxa: int = _CONTENT_WIDTH_DXA
    ) -> List[int]:
        cols = max((len(r) for r in values), default=1)
        weights: List[float] = []
        for c in range(cols):
            maxlen = 4.0
            for row in values:
                if c >= len(row):
                    continue
                s = str(row[c] or "")
                w = sum(2.0 if ord(ch) > 127 else 1.0 for ch in s)
                maxlen = max(maxlen, w)
            weights.append(maxlen)
        total_w = sum(weights) or 1.0
        raw = [max(0.08, w / total_w) for w in weights]
        scale = 1.0 / sum(raw)
        widths = [max(400, int(total_dxa * r * scale)) for r in raw]
        drift = total_dxa - sum(widths)
        if widths:
            widths[-1] = max(400, widths[-1] + drift)
        return widths

    # ------------------------------------------------------------------ cover tables

    def _tighten_cover_info_table(self, doc: Document) -> None:
        """Keep template 委托单位/地址 rows at golden cover height."""
        for table in doc.tables:
            texts = [c.text for row in table.rows for c in row.cells]
            joined = "".join(texts)
            if "委托单位" not in joined:
                continue
            self._set_table_cell_margins(table, top=0, left=0, bottom=0, right=0)
            for row in table.rows:
                self._set_row_height(row, COVER_ROW_HEIGHT_TWIPS)
                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for p in cell.paragraphs:
                        self._apply_tight_spacing(p)
            break

    def _insert_sample_info_table(
        self, doc: Document, anchor: Paragraph, state: ProjectState
    ):
        lang = self._lang()
        side = self._side_lang()
        rows_data: List[Tuple[str, str]] = []
        for key, val in state.iter_overview_fields(side):
            if key in _COVER_SKIP:
                continue
            if ProjectState.is_custom_overview_key(key):
                label = state.overview_display_label(key, side)
                if not label and not (val or "").strip():
                    continue
            else:
                label = field_label(key, lang)
            if lang == "英文" and not label and not ProjectState.is_custom_overview_key(key):
                rows_data.append(("", val))
                continue
            rows_data.append((label or key, val))

        qty = (state.overview_field_map(side).get("送样数量") or "").strip()
        if qty and not qty.lower().endswith("pcs"):
            qty = f"{qty}pcs"
        if qty:
            rows_data.append((field_label("客户送样数量", lang) or "客户送样数量", qty))

        recv = self._fmt_date(state.sample_receive_date)
        if recv:
            rows_data.append((field_label("样品接收日期", lang) or "样品接收日期", recv))

        period = self._fmt_period(state.test_start_date, state.test_end_date)
        if period:
            rows_data.append((field_label("样品检测日期", lang) or "样品检测日期", period))

        table = self._add_table_before(
            doc, anchor, max(len(rows_data), 1), 2, bordered=False
        )
        self._set_col_widths(table, _WIDTHS_COVER_INFO)
        self._set_table_cell_margins(table, top=0, left=40, bottom=0, right=40)
        for i, (label, value) in enumerate(rows_data):
            self._set_row_height(table.rows[i], COVER_ROW_HEIGHT_TWIPS)
            self._set_cell_text(
                table.rows[i].cells[0], label, align=WD_ALIGN_PARAGRAPH.LEFT
            )
            self._set_cell_text(
                table.rows[i].cells[1],
                f"：{value}" if value else "：",
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )

        note_zh = "以上测试之样品及信息是由申请者提供并确认"
        note_en = (
            "The above sample(s) and information are provided and confirmed by the applicant"
        )
        if lang == "中英文":
            self._add_para_before(doc, anchor, note_zh, size=SIZE_BODY).alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
            )
            self._add_para_before(doc, anchor, note_en, size=SIZE_BODY).alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
            )
        else:
            note_p = self._add_para_before(
                doc, anchor, note_en if lang == "英文" else note_zh, size=SIZE_BODY
            )
            note_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _insert_sample_list_table(
        self, doc: Document, anchor: Paragraph, state: ProjectState, nodes: List[TestNode]
    ):
        lang = self._lang()
        side = self._side_lang()
        fields = state.overview_field_map(side)
        if lang == "英文":
            sample_name = fields.get("样品名称") or state.sample_name_en or "/"
            part_no = fields.get("零件号") or "/"
        else:
            sample_name = fields.get("样品名称") or state.sample_name or "/"
            part_no = fields.get("零件号") or "/"
            if lang == "中英文":
                sample_name = language_text(
                    sample_name if sample_name != "/" else (state.sample_name or ""),
                    (state.application_fields_en or {}).get("样品名称")
                    or state.sample_name_en
                    or "",
                    "中英文",
                ) or "/"
                part_no = language_text(
                    part_no if part_no != "/" else "",
                    (state.application_fields_en or {}).get("零件号") or "",
                    "中英文",
                ) or "/"
        qty = (fields.get("送样数量") or "").strip() or str(
            max((len(n.samples) for n in nodes), default=0) or ""
        )
        if qty and not str(qty).lower().endswith("pcs"):
            qty = f"{qty}pcs"

        ids = self._collect_sample_ids(nodes)
        id_text = self._format_id_range(ids) if ids else "/"

        table = self._add_table_before(doc, anchor, 2, 5)
        self._set_col_widths(table, _WIDTHS_SAMPLE_LIST)
        headers = [
            self._H("序号", "No."),
            self._H("样品名称", "Sample Name"),
            self._H("零件号", "Part No."),
            self._H("送样数量", "Quantity of Samples"),
            self._H("样品编号", "Sample No."),
        ]
        for i, h in enumerate(headers):
            self._set_cell_text(table.rows[0].cells[i], h)
        values = ["1", sample_name, part_no, qty, id_text]
        for i, v in enumerate(values):
            self._set_cell_text(table.rows[1].cells[i], v)

    def _insert_summary_table(self, doc: Document, anchor: Paragraph, nodes: List[TestNode]):
        table = self._add_table_before(doc, anchor, max(len(nodes), 0) + 1, 6)
        self._set_col_widths(table, _WIDTHS_SUMMARY)
        headers = [
            self._H("序号", "No."),
            self._H("检测项目", "Test Item"),
            self._H("检测方法", "Test Method"),
            self._H("检测时间", "Testing Period"),
            self._H("样品编号", "Sample No."),
            self._H("结论", "Conclusion"),
        ]
        for i, h in enumerate(headers):
            self._set_cell_text(table.rows[0].cells[i], h)
        for idx, node in enumerate(nodes, 1):
            ids = [s.sample_id for s in node.samples if s.sample_id]
            conclusion = self._node_conclusion(node)
            vals = [
                str(idx),
                self._test_item_label(node) or "/",
                self._format_method(node) or "/",
                self._fmt_period(node.start_date, node.end_date) or "/",
                self._format_id_range(ids) if ids else "/",
                conclusion,
            ]
            for i, v in enumerate(vals):
                self._set_cell_text(table.rows[idx].cells[i], v)

    def _env_condition_text(self, node: TestNode) -> str:
        text = node.resolved_env_condition()
        return text or ENV_CONDITION_TEXT

    def _test_item_label(self, node: TestNode) -> str:
        zh = (node.test_name or "").strip()
        en = (node.test_name_en or "").strip()
        return language_text(zh, en, self._lang())

    @staticmethod
    def _collect_sample_ids(nodes: Iterable[TestNode]) -> List[str]:
        seen = set()
        out: List[str] = []
        for node in nodes:
            for s in node.samples or []:
                sid = (s.sample_id or "").strip()
                if sid and sid not in seen:
                    seen.add(sid)
                    out.append(sid)
        return out

    @staticmethod
    def _format_id_range(ids: Sequence[str]) -> str:
        clean = [i.strip() for i in ids if i and str(i).strip()]
        if not clean:
            return "/"
        if len(clean) == 1:
            return clean[0]
        # TP-262686912 ~ TP-262686919 style when sharing a numeric suffix
        def split_tail(s: str):
            m = re.match(r"^(.*?)(\d+)$", s)
            return (m.group(1), m.group(2)) if m else (s, "")

        prefix0, num0 = split_tail(clean[0])
        prefix1, num1 = split_tail(clean[-1])
        if prefix0 and prefix0 == prefix1 and num0 and num1 and len(clean) > 1:
            return f"{clean[0]}~{clean[-1]}"
        return "、".join(clean)

    @staticmethod
    def _format_method(node: TestNode) -> str:
        parts = []
        for std in node.resolved_standards():
            sid = (std.standard_id or "").strip()
            chap = (std.chapter or "").strip()
            if sid and chap:
                parts.append(f"{sid}-{chap}")
            elif sid:
                parts.append(sid)
            elif chap:
                parts.append(chap)
        if parts:
            return "；".join(parts)
        return (node.joined_test_method() or "").replace(" / ", "-")

    def _node_conclusion(self, node: TestNode) -> str:
        results = []
        for s in (node.samples or []):
            if not s.sample_id:
                continue
            results.extend(s.all_results())
        if not results:
            return format_conclusion(TestResult.NA, self._lang())
        if any(r == TestResult.FAIL for r in results):
            return format_conclusion(TestResult.FAIL, self._lang())
        if all(r == TestResult.PASS for r in results):
            return format_conclusion(TestResult.PASS, self._lang())
        if all(r == TestResult.NA for r in results):
            return format_conclusion(TestResult.NA, self._lang())
        if any(r == TestResult.FAIL for r in results):
            return format_conclusion(TestResult.FAIL, self._lang())
        return format_conclusion(TestResult.PASS, self._lang())

    # ------------------------------------------------------------------ test details

    def _insert_test_details(
        self,
        doc: Document,
        anchor: Paragraph,
        nodes: List[Tuple[TestLeg, TestNode]],
        project_path: Optional[str],
    ):
        for idx, (leg, node) in enumerate(nodes, 1):
            self._append_test_node(doc, anchor, leg, node, idx, project_path)

    def _append_test_node(
        self,
        doc: Document,
        anchor: Paragraph,
        leg: TestLeg,
        node: TestNode,
        index: int,
        project_path: Optional[str],
    ):
        # Each test item starts on a new page (matches CTI golden layout)
        self._add_page_break_before(doc, anchor)
        item = self._test_item_label(node) or "/"
        self._add_para_before(
            doc,
            anchor,
            f"{index}. {self._L('检测项目', 'Test Item')}：{item}",
            size=SIZE_BODY,
        )
        self._add_para_before(
            doc,
            anchor,
            f"（1）{self._L('检测环境条件', 'Test environment')}：{self._env_condition_text(node)}",
            size=SIZE_BODY,
        )

        self._add_para_before(
            doc, anchor, f"（2）{self._L('检测设备', 'Test equipment')}：", size=SIZE_BODY
        )
        self._insert_equipment_table(doc, anchor, node)

        method = self._format_method(node) or "/"
        self._add_para_before(
            doc,
            anchor,
            f"（3）{self._L('检测方法', 'Test method')}：{method}",
            size=SIZE_BODY,
        )

        self._add_para_before(
            doc, anchor, f"（4）{self._L('检测条件', 'Test condition')}：", size=SIZE_BODY
        )
        self._insert_condition_blocks(doc, anchor, node)
        self._insert_condition_images(doc, anchor, node)

        qty = len([s for s in (node.samples or []) if s.sample_id])
        self._add_para_before(
            doc,
            anchor,
            f"（5）{self._L('检测数量', 'Quantity')}：{qty}pcs",
            size=SIZE_BODY,
        )

        self._add_para_before(
            doc,
            anchor,
            f"（6）{self._L('评判要求', 'Evaluation requirement')}：",
            size=SIZE_BODY,
        )
        eval_req = self._evaluation_text(node)
        if eval_req:
            for block in re.split(r"\n+", eval_req):
                if block.strip():
                    self._add_para_before(doc, anchor, block.strip(), size=SIZE_BODY)

        self._add_para_before(
            doc, anchor, f"（7）{self._L('检测结果', 'Test result')}：", size=SIZE_BODY
        )
        self._insert_sample_result_table(doc, anchor, node)
        self._insert_data_tables(doc, anchor, node, project_path, leg.leg_name)

        conclusion = self._node_conclusion(node)
        self._add_para_before(
            doc,
            anchor,
            f"（8）{self._L('结论', 'Conclusion')}：{conclusion}",
            size=SIZE_BODY,
        )

        # Photos section on its own page
        self._add_page_break_before(doc, anchor)
        self._add_para_before(
            doc,
            anchor,
            self._L("检测照片", "Test photos"),
            size=SIZE_PHOTO_TITLE,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        if project_path and node.test_name:
            self._insert_photos(
                doc,
                anchor,
                Path(project_path),
                leg.leg_name,
                node.test_name,
                order=getattr(node, "photo_album_order", None) or None,
            )

    def _evaluation_text(self, node: TestNode) -> str:
        lang = self._lang()
        if lang == "英文":
            return (node.evaluation_req_en or node.joined_evaluation_req_en() or "").strip()
        if lang == "中英文":
            zh = (node.evaluation_req or node.joined_evaluation_req() or "").strip()
            en = (node.evaluation_req_en or node.joined_evaluation_req_en() or "").strip()
            return language_text(zh, en, "中英文") or "/"
        return (node.evaluation_req or node.joined_evaluation_req() or "").strip() or "/"

    def _insert_equipment_table(self, doc: Document, anchor: Paragraph, node: TestNode):
        items = list(node.equipments or [])
        headers = [
            self._H("序号", "No."),
            self._H("名称", "Name"),
            self._H("型号", "Model"),
            self._H("设备编号", "Equipment No."),
            self._H("校准有效期", "Calibration due date"),
        ]
        if not items and node.equipment_name:
            table = self._add_table_before(doc, anchor, 2, 5)
            self._set_col_widths(table, _WIDTHS_EQUIPMENT)
            for i, h in enumerate(headers):
                self._set_cell_text(table.rows[0].cells[i], h)
            self._set_cell_text(table.rows[1].cells[0], "1")
            name = node.equipment_name if self._lang() != "英文" else ""
            self._set_cell_text(table.rows[1].cells[1], name or "/")
            for i in range(2, 5):
                self._set_cell_text(table.rows[1].cells[i], "/")
            return
        if not items:
            self._add_para_before(doc, anchor, "/", size=SIZE_BODY)
            return

        table = self._add_table_before(doc, anchor, len(items) + 1, 5)
        self._set_col_widths(table, _WIDTHS_EQUIPMENT)
        for i, h in enumerate(headers):
            self._set_cell_text(table.rows[0].cells[i], h)
        for idx, eq in enumerate(items, 1):
            name = language_text(
                (eq.name or "").replace("\n", ""),
                (getattr(eq, "name_en", None) or "").replace("\n", ""),
                self._lang(),
            ) or "/"
            model = eq.model or "/"
            code = eq.code or "/"
            cal = (getattr(eq, "valid_date", None) or "").strip() or "/"
            vals = [str(idx), name, model, code, cal]
            for i, v in enumerate(vals):
                self._set_cell_text(table.rows[idx].cells[i], v)

    def _insert_condition_blocks(self, doc: Document, anchor: Paragraph, node: TestNode):
        """检测条件：每条标准先写试验名称，再写缩进两格的条件正文。"""
        stds = node.resolved_standards()
        lang = self._lang()
        indent = "  "

        def emit_body(zh: str, en: str):
            zh = (zh or "").strip()
            en = (en or "").strip()
            if lang == "英文":
                texts = [en] if en else [""]
            elif lang == "中英文":
                texts = [t for t in (zh, en) if t] or ["/"]
            else:
                texts = [zh] if zh else ["/"]
            for body in texts:
                if not body and lang == "英文":
                    continue
                for block in re.split(r"\n+", body):
                    text = block.strip()
                    if text:
                        self._add_para_before(doc, anchor, indent + text, size=SIZE_BODY)

        if not stds:
            emit_body(node.standard_desc or "", node.standard_desc_en or "")
            return
        for std in stds:
            title = (std.field_title() or std.test_name or "").strip() or "/"
            self._add_para_before(doc, anchor, title, size=SIZE_BODY)
            emit_body(std.standard_desc or "", std.standard_desc_en or "")

    def _insert_condition_images(self, doc: Document, anchor: Paragraph, node: TestNode):
        blobs: List[bytes] = []
        for std in node.resolved_standards():
            for blob in std.images or []:
                if blob:
                    blobs.append(blob)
        for blob in blobs:
            width = self._condition_image_width_in(blob)
            self._add_picture_bytes_before(doc, anchor, blob, width)

    def _result_desc_texts(self, node: TestNode) -> List[str]:
        """One 试验结果 text per result table, in selection order.

        Multi-standard → one entry per standard that still has a result table.
        Single / legacy → sample or node fallback handled by the caller per row.
        """
        lang = self._lang()
        stds = node.result_table_standards()

        def side_text(std) -> str:
            zh = (std.result_desc or "").strip()
            en = (std.result_desc_en or "").strip()
            if lang == "英文":
                return en
            if lang == "中英文":
                return language_text(zh, en, "中英文")
            return zh

        if len(stds) > 1:
            return [side_text(s) or ("/" if lang != "英文" else "") for s in stds]
        if len(stds) == 1:
            text = side_text(stds[0])
            return [text] if text else [""]
        return [""]

    def _insert_sample_result_table(self, doc: Document, anchor: Paragraph, node: TestNode):
        """One 样品编号/试验结果/试验结论 table per visible result-table standard."""
        samples = [s for s in (node.samples or []) if s.sample_id]
        stds = node.result_table_standards()
        if not stds:
            if not samples:
                self._add_para_before(
                    doc, anchor, self._L("无结果记录", "No result recorded"), size=SIZE_BODY
                )
            return
        if not samples:
            self._add_para_before(
                doc, anchor, self._L("无结果记录", "No result recorded"), size=SIZE_BODY
            )
            return
        descs = self._result_desc_texts(node)
        lang = self._lang()
        node_desc = language_text(
            (getattr(node, "result_desc", None) or "").strip(),
            (getattr(node, "result_desc_en", None) or "").strip(),
            lang,
        )
        multi = len(descs) > 1
        headers = [
            self._H("样品编号", "Sample No."),
            self._H("试验结果", "Test result"),
            self._H("试验结论", "Conclusion"),
        ]
        for ti, table_desc in enumerate(descs):
            if ti:
                self._add_para_before(doc, anchor, "", size=SIZE_BODY)
            table = self._add_table_before(doc, anchor, len(samples) + 1, 3)
            self._set_col_widths(table, _WIDTHS_SAMPLE_RESULT)
            for i, h in enumerate(headers):
                self._set_cell_text(table.rows[0].cells[i], h)
            self._set_row_as_tbl_header(table.rows[0])
            std = stds[ti] if ti < len(stds) else None
            for idx, sample in enumerate(samples, 1):
                if multi:
                    desc = table_desc
                else:
                    sample_zh = (getattr(sample, "result_desc", None) or "").strip()
                    if lang == "英文":
                        desc = table_desc or node_desc or ""
                    else:
                        desc = sample_zh or table_desc or node_desc or "/"
                if std is not None:
                    conclusion = format_conclusion(sample.result_for(std), lang)
                else:
                    conclusion = format_conclusion(sample.result, lang)
                vals = [
                    sample.sample_id,
                    desc,
                    conclusion,
                ]
                for i, v in enumerate(vals):
                    self._set_cell_text(table.rows[idx].cells[i], v)

    def _insert_data_tables(
        self,
        doc: Document,
        anchor: Paragraph,
        node: TestNode,
        project_path: Optional[str],
        leg_name: str = "",
    ):
        root = Path(project_path) if project_path else None
        refs: List[DataTableRef] = []
        if root is not None and leg_name and node.test_name:
            refs = list_attachment_refs(root, leg_name, node.test_name)
        if not refs:
            refs = list(getattr(node, "data_tables", None) or [])
        if not refs or root is None:
            return
        self._add_para_before(
            doc,
            anchor,
            self._L("试验数据", "Test data"),
            size=SIZE_BODY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        for ref in refs:
            try:
                path = resolve_attachment_path(root, ref)
                snap = read_preview_snapshot(path)
            except Exception:
                continue
            if not snap.values:
                continue
            title = Path(ref.title or path.stem).stem.strip()
            if title:
                self._add_para_before(
                    doc, anchor, title, size=SIZE_BODY, align=WD_ALIGN_PARAGRAPH.CENTER
                )
            rows = len(snap.values)
            cols = max((len(r) for r in snap.values), default=1)
            table = self._add_table_before(doc, anchor, rows, cols)
            self._set_col_widths(table, self._estimate_col_widths(snap.values))
            self._set_table_cell_margins(table, top=20, left=40, bottom=20, right=40)
            # Merge first, then write only into merge anchors — avoids empty ¶ stacking
            self._apply_snapshot_merges(table, snap)
            slaves = self._merged_slave_cells(snap)
            for r_i, row in enumerate(snap.values):
                for c_i in range(cols):
                    if (r_i, c_i) in slaves:
                        continue
                    val = row[c_i] if c_i < len(row) else ""
                    self._set_cell_text(table.rows[r_i].cells[c_i], val)
            n_header = infer_header_row_count(snap)
            for i in range(min(n_header, rows)):
                self._set_row_as_tbl_header(table.rows[i])

    def _insert_photos(
        self,
        doc: Document,
        anchor: Paragraph,
        project_root: Path,
        leg_name: str,
        test_name: str,
        order=None,
    ):
        albums = list_albums(project_root, leg_name, test_name, order=order)
        if not albums:
            return
        lang = self._lang()
        for album in albums:
            photos = list_photos(project_root, leg_name, test_name, album)
            if not photos:
                continue
            if uses_data_photo_layout(album):
                for path in photos:
                    self._add_picture_file_before(doc, anchor, path, DATA_PHOTO_WIDTH_IN)
                    cap_text = photo_caption(album, lang, file_stem=path.stem)
                    cap = self._add_para_before(doc, anchor, cap_text, size=SIZE_CAPTION)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                self._insert_photo_pairs(doc, anchor, photos, PHOTO_WIDTH_IN, album=album)

    def _insert_photo_pairs(
        self,
        doc: Document,
        anchor: Paragraph,
        photos: Sequence[Path],
        width_in: float,
        album: str = "",
    ):
        pairs = list(photos)
        half = _CONTENT_WIDTH_DXA // 2
        lang = self._lang()
        for i in range(0, len(pairs), 2):
            chunk = pairs[i : i + 2]
            table = self._add_table_before(doc, anchor, 1, 2, bordered=False)
            self._set_col_widths(table, (half, half))
            for col, path in enumerate(chunk):
                cell = table.rows[0].cells[col]
                cell.text = ""
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._apply_tight_spacing(p)
                try:
                    stream = self._prepare_embed_stream(path, width_in)
                    run = p.add_run()
                    run.add_picture(stream, width=Inches(width_in))
                except Exception as exc:
                    err = p.add_run(f"[无法插入图片: {path.name}]")
                    self._style_run(err, size=SIZE_CAPTION)
                    print(f"Failed to add image {path}: {exc}")
                cap_text = photo_caption(album, lang, file_stem=path.stem)
                cap = cell.add_paragraph(cap_text)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._apply_tight_spacing(cap)
                if cap.runs:
                    self._style_run(cap.runs[0], size=SIZE_CAPTION)
            if len(chunk) == 1:
                table.rows[0].cells[1].text = ""

    def _add_picture_file_before(
        self, doc: Document, anchor: Paragraph, path: Path, width_in: float
    ):
        p = self._add_para_before(doc, anchor, "")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            stream = self._prepare_embed_stream(path, width_in)
            run = p.add_run()
            run.add_picture(stream, width=Inches(width_in))
        except Exception as exc:
            p.add_run(f"[无法插入图片: {path.name}]")
            print(f"Failed to add image {path}: {exc}")

    def _add_picture_bytes_before(
        self, doc: Document, anchor: Paragraph, blob: bytes, width_in: float
    ):
        """Embed standard-library condition images as-is (no resize/re-encode)."""
        p = self._add_para_before(doc, anchor, "")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = p.add_run()
            run.add_picture(io.BytesIO(blob), width=Inches(width_in))
        except Exception as exc:
            p.add_run("[无法插入条件图片]")
            print(f"Failed to add condition image: {exc}")
