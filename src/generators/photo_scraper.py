from pathlib import Path
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from src.io.project_assets import iter_merged_export_photos
from src.io.test_photos import uses_data_photo_layout


class PhotoScraper:
    """Legacy helper; WordGenerator embeds photos directly with 2-up layout."""

    def __init__(self, base_project_dir: str, remote_root: str = None):
        self.base_project_dir = Path(base_project_dir)
        self.remote_root = Path(remote_root) if remote_root else None

    def add_photos_to_document(
        self,
        doc: Document,
        leg_name: str,
        test_name: str,
        max_width_inches: float = 2.95,
        data_width_inches: float = 5.5,
        order=None,
    ):
        temps = []
        try:
            exported = iter_merged_export_photos(
                self.base_project_dir,
                self.remote_root,
                leg_name,
                test_name,
                order=order,
                temps=temps,
            )
            if not exported:
                return
            doc.add_paragraph("检测样品照片:", style="Normal")
            by_album = []
            for item in exported:
                if not by_album or by_album[-1][0] != item.album:
                    by_album.append((item.album, [item]))
                else:
                    by_album[-1][1].append(item)
            for album, items in by_album:
                if uses_data_photo_layout(album):
                    for item in items:
                        self._add_one(doc, item.path, data_width_inches, stem=item.stem)
                else:
                    paths = [item.path for item in items]
                    for i in range(0, len(paths), 2):
                        chunk = paths[i : i + 2]
                        table = doc.add_table(rows=1, cols=2)
                        for col, img_path in enumerate(chunk):
                            cell = table.rows[0].cells[col]
                            cell.text = ""
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            try:
                                with Image.open(img_path) as img:
                                    img.size
                                p.add_run().add_picture(
                                    str(img_path), width=Inches(max_width_inches)
                                )
                            except Exception as e:
                                print(f"Failed to add image {img_path}: {e}")
                            stem = items[i + col].stem
                            cap = cell.add_paragraph(stem)
                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        finally:
            for temp in temps:
                try:
                    Path(temp).unlink(missing_ok=True)
                except OSError:
                    pass

    def _add_one(self, doc: Document, img_path: Path, width: float, stem: str = ""):
        try:
            with Image.open(img_path) as img:
                img.size
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(img_path), width=Inches(width))
            caption = doc.add_paragraph(stem or img_path.stem)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Failed to add image {img_path}: {e}")
